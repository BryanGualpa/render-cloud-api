"""Genera el documento Word de presentación AnalytiCore."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(11)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def build_document():
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # PORTADA
    doc.add_paragraph()
    doc.add_paragraph()
    add_title(doc, 'AnalytiCore')
    add_subtitle(doc, 'Prototipo de Arquitectura Orientada a Servicios en la Nube')
    doc.add_paragraph()
    add_subtitle(doc, 'Plataforma de Análisis de Sentimiento y Palabras Clave')
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Arquitectura de Software / Cloud Computing\nRender — Oregon (US West)')
    r.font.size = Pt(12)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Repositorio: https://github.com/BryanGualpa/render-cloud-api')
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    doc.add_page_break()

    # 1. CONTEXTO
    add_heading(doc, '1. Contexto del Negocio')
    add_para(doc,
        'AnalytiCore es una startup de análisis de datos que desea ofrecer un servicio en línea '
        'para que los usuarios envíen textos y obtengan un análisis simple de sentimiento y '
        'extracción de palabras clave. El objetivo del proyecto es construir un prototipo funcional '
        'desplegado en la nube que demuestre la viabilidad técnica de una arquitectura políglota '
        '(React, Python y Java) y sirva como base para futuras expansiones comerciales.')
    add_para(doc,
        'El reto principal consiste en validar que es posible procesar solicitudes de usuarios de '
        'forma confiable, mantener el estado de los trabajos de análisis y escalar cada componente '
        'de forma independiente, sin invertir en infraestructura propia.')

    # 2. SOLUCIÓN
    add_heading(doc, '2. Solución Propuesta')
    add_para(doc,
        'Se desarrolló un prototipo con arquitectura orientada a servicios (SOA) desplegado en '
        'la plataforma Render, compuesto por cuatro componentes principales:')
    add_bullet(doc, 'Frontend Web (React + Nginx): interfaz de usuario para enviar textos y consultar resultados.')
    add_bullet(doc, 'Servicio de Submisión (Python/Flask): recibe solicitudes, valida datos, persiste trabajos y orquesta el análisis.')
    add_bullet(doc, 'Servicio de Análisis (Java/Spring Boot): procesa el texto, calcula sentimiento y extrae palabras clave.')
    add_bullet(doc, 'Base de Datos PostgreSQL (Render Managed): almacena el estado y los resultados de cada trabajo.')
    add_para(doc,
        'La comunicación entre componentes es exclusivamente mediante APIs RESTful en formato JSON. '
        'Ningún servicio guarda estado en memoria; toda la información se externaliza a PostgreSQL.')

    # 3. ARQUITECTURA
    add_heading(doc, '3. Visión General de la Arquitectura')
    add_para(doc, 'La arquitectura desplegada en Render (región Oregon) sigue el siguiente flujo:')
    flow_steps = [
        ('1', 'Usuario → Frontend', 'El usuario introduce el texto en la interfaz React y lo envía.'),
        ('2', 'Frontend → Python', 'POST /api/jobs — el servicio Python crea un registro PENDIENTE en PostgreSQL.'),
        ('3', 'Python → Java', 'POST /api/analyze — notifica al servicio Java que hay un trabajo listo.'),
        ('4', 'Java ↔ PostgreSQL', 'Actualiza estado a PROCESANDO, analiza el texto y guarda resultados como COMPLETADO.'),
        ('5', 'Frontend → Python', 'Polling GET /api/jobs/{jobId} cada 2 segundos hasta obtener los resultados.'),
    ]
    add_table(doc, ['Paso', 'Comunicación', 'Descripción'], flow_steps)

    # 4. COMPONENTES
    add_heading(doc, '4. Componentes del Sistema')
    add_table(doc,
        ['Componente', 'Tecnología', 'Responsabilidad', 'Docker'],
        [
            ('Frontend', 'React + Vite + Nginx', 'Interfaz de usuario, envío y consulta de resultados', 'Sí'),
            ('Python Service', 'Flask + Gunicorn', 'Submisión, validación, persistencia y orquestación', 'Sí'),
            ('Java Service', 'Spring Boot 3.2', 'Análisis de sentimiento y extracción de keywords', 'Sí'),
            ('PostgreSQL', 'Render Managed DB', 'Persistencia de trabajos (analysis_jobs)', 'Gestionado'),
        ])

    # 5. PATRONES CLOUD
    add_heading(doc, '5. Patrones de Arquitectura Cloud Aplicados')
    patterns = [
        ('Empaquetado Docker', 'Cada componente (Frontend, Python, Java) está en su propia imagen Docker para producción.'),
        ('APIs RESTful', 'Toda la comunicación entre servicios es mediante HTTP/JSON.'),
        ('Servicios Stateless', 'Ningún servicio guarda estado en memoria; PostgreSQL centraliza los datos.'),
        ('Configuración externa', 'DATABASE_URL y JAVA_SERVICE_URL se inyectan como variables de entorno.'),
        ('Arquitectura limpia', 'Cada servicio tiene capas separadas: Presentación, Aplicación, Dominio e Infraestructura.'),
    ]
    add_table(doc, ['Patrón', 'Implementación'], patterns)

    # 6. ARQUITECTURA LIMPIA
    add_heading(doc, '6. Arquitectura Limpia por Componente')
    add_heading(doc, '6.1 Frontend (React)', level=2)
    add_bullet(doc, 'Presentación: componentes React (App.jsx), estilos CSS.')
    add_bullet(doc, 'Aplicación: lógica de envío de formulario y polling de resultados.')
    add_bullet(doc, 'Infraestructura: Nginx como servidor web y proxy hacia la API Python.')
    add_heading(doc, '6.2 Python Service (Submisión)', level=2)
    add_bullet(doc, 'Presentación: Blueprint Flask con rutas /api/jobs y /health.')
    add_bullet(doc, 'Aplicación: CreateJobUseCase, GetJobUseCase.')
    add_bullet(doc, 'Dominio: entidad Job, interfaz JobRepository.')
    add_bullet(doc, 'Infraestructura: PostgresJobRepository, JavaAnalysisClient, connection.py.')
    add_heading(doc, '6.3 Java Service (Análisis)', level=2)
    add_bullet(doc, 'Presentación: AnalysisController, HealthController (REST).')
    add_bullet(doc, 'Aplicación: AnalysisService (lógica de sentimiento y keywords).')
    add_bullet(doc, 'Dominio: Job, AnalysisResult, JobRepository (interfaz).')
    add_bullet(doc, 'Infraestructura: PostgresJobRepository, DataSourceConfig.')

    # 7. DESPLIEGUE
    add_heading(doc, '7. Despliegue en Render')
    add_table(doc,
        ['Servicio Render', 'Root Directory', 'Runtime', 'Variables de entorno'],
        [
            ('render-cloud-db', '—', 'PostgreSQL Free', 'Internal Database URL'),
            ('analyticore-java', 'java-service', 'Docker', 'DATABASE_URL'),
            ('render-cloud-api', 'python-service', 'Docker', 'DATABASE_URL, JAVA_SERVICE_URL'),
            ('analyticore-frontend', 'frontend', 'Docker', 'API_BACKEND'),
        ])

    # 8. DEMO
    add_heading(doc, '8. Demostración Funcional (URLs en Vivo)')
    add_para(doc, 'El prototipo está desplegado y operativo en los siguientes enlaces:', bold=True)
    demo_links = [
        ('Aplicación Web (Demo principal)', 'https://analyticore-frontend-p4ab.onrender.com'),
        ('API Python — Health Check', 'https://render-cloud-api-tyio.onrender.com/health'),
        ('API Java — Health Check', 'https://analyticore-java-k3yk.onrender.com/health'),
        ('Repositorio GitHub', 'https://github.com/BryanGualpa/render-cloud-api'),
    ]
    add_table(doc, ['Recurso', 'URL'], demo_links)
    add_para(doc, 'Prueba de funcionamiento verificada:')
    add_bullet(doc, 'Texto enviado: "Me encanta este proyecto, es excelente y muy util"')
    add_bullet(doc, 'Resultado: Sentimiento POSITIVO, estado COMPLETADO, palabras clave extraídas.')
    add_bullet(doc, 'Health checks de Python y Java responden correctamente con status healthy.')

    # 9. ENDPOINTS
    add_heading(doc, '9. Endpoints de la API')
    add_table(doc,
        ['Servicio', 'Método', 'Endpoint', 'Descripción'],
        [
            ('Python', 'GET', '/health', 'Verifica que el servicio de submisión está activo'),
            ('Python', 'POST', '/api/jobs', 'Envía texto para análisis, retorna jobId'),
            ('Python', 'GET', '/api/jobs/{id}', 'Consulta estado y resultados de un trabajo'),
            ('Java', 'GET', '/health', 'Verifica que el servicio de análisis está activo'),
            ('Java', 'POST', '/api/analyze', 'Inicia análisis de un jobId (uso interno)'),
        ])

    # 10. BENEFICIOS
    add_heading(doc, '10. Beneficios de la Arquitectura Elegida')
    benefits = [
        ('Escalabilidad', 'Cada servicio escala de forma independiente según su carga.'),
        ('Mantenibilidad', 'Capas separadas facilitan cambios sin afectar otros componentes.'),
        ('Flexibilidad del equipo', 'Tecnologías políglotas permiten equipos especializados en paralelo.'),
        ('Resiliencia', 'Servicios stateless + BD centralizada = recuperación ante fallos.'),
        ('Costo', 'Plan Free de Render permite validar el concepto sin inversión inicial.'),
    ]
    add_table(doc, ['Beneficio', 'Descripción'], benefits)

    # 11. ENTREGABLES
    add_heading(doc, '11. Entregables del Proyecto')
    deliverables = [
        ('Código fuente', '/frontend, /python-service, /java-service con Dockerfile en cada uno'),
        ('Diagrama de componentes', 'docs/diagrama-componentes.md'),
        ('Diagramas de capas (×3)', 'docs/capas-frontend.md, capas-python-service.md, capas-java-service.md'),
        ('Informe ejecutivo', 'docs/informe-ejecutivo.md (máximo 1 página)'),
        ('Repositorio GitHub', 'https://github.com/BryanGualpa/render-cloud-api'),
        ('Demo en vivo', 'https://analyticore-frontend-p4ab.onrender.com'),
    ]
    add_table(doc, ['Entregable', 'Ubicación'], deliverables)

    # 12. CONCLUSIÓN
    add_heading(doc, '12. Conclusión')
    add_para(doc,
        'AnalytiCore demuestra que una arquitectura cloud orientada a servicios, con comunicación REST, '
        'persistencia externalizada en PostgreSQL y despliegue en contenedores Docker, es una base sólida '
        'y escalable para convertir el análisis de texto en un producto comercial. El prototipo cumple '
        'con todos los requisitos técnicos y funcionales: arquitectura políglota, servicios sin estado, '
        'APIs RESTful, arquitectura limpia en cada componente y demostración funcional en la nube.')
    add_para(doc,
        'La solución permite a AnalytiCore validar su modelo de negocio con costo mínimo (plan Free de Render) '
        'y prepara el camino para futuras expansiones como autenticación de usuarios, análisis avanzado con IA '
        'y escalado horizontal de los servicios de análisis.')

    # GUÍA DE PRESENTACIÓN
    doc.add_page_break()
    add_heading(doc, 'Anexo: Guía para Presentar al Profesor (5 minutos)')
    add_para(doc, 'Sigue estos pasos durante la exposición:', bold=True)
    steps = [
        'Mostrar el repositorio GitHub con la estructura /frontend, /python-service, /java-service.',
        'Abrir docs/diagrama-componentes.md y explicar el flujo de datos entre los 4 componentes.',
        'Abrir https://analyticore-frontend-p4ab.onrender.com y escribir un texto de ejemplo.',
        'Hacer clic en "Enviar análisis" y mostrar el sentimiento y las palabras clave.',
        'Abrir /health del Python y Java en pestañas para demostrar que los microservicios responden.',
        'Mencionar los patrones cloud: Docker, REST, Stateless, configuración externa, arquitectura limpia.',
    ]
    for i, step in enumerate(steps, 1):
        add_bullet(doc, f'{i}. {step}')

    return doc


if __name__ == '__main__':
    output = Path(__file__).parent / 'AnalytiCore_Presentacion.docx'
    build_document().save(output)
    print(f'Documento generado: {output}')
